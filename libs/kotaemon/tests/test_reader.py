from pathlib import Path
from unittest.mock import patch

from langchain_core.documents import Document as LangchainDocument
from llama_index.core.node_parser import SimpleNodeParser

from kotaemon.base import Document
from kotaemon.loaders import (
    AutoReader,
    AzureAIDocumentIntelligenceLoader,
    DocxReader,
    HtmlReader,
    MhtmlReader,
    UnstructuredReader,
)

from .conftest import skip_when_unstructured_pdf_not_installed


def test_docx_reader():
    reader = DocxReader()
    documents = reader.load_data(Path(__file__).parent / "resources" / "dummy.docx")

    assert len(documents)


def test_docx_reader_preserves_table_cell_text(tmp_path):
    from docx import Document as DocxDocument

    docx_path = tmp_path / "table-answer.docx"
    document = DocxDocument()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Owner"
    table.cell(1, 1).text = "Task 8 Word table answer"
    document.save(docx_path)

    documents = DocxReader().load_data(docx_path)

    assert any("Task 8 Word table answer" in doc.text for doc in documents)


def test_default_file_extractors_route_docx_to_table_preserving_reader():
    from kotaemon.indices.ingests.files import KH_DEFAULT_FILE_EXTRACTORS

    assert isinstance(KH_DEFAULT_FILE_EXTRACTORS[".docx"], DocxReader)


def test_html_reader():
    reader = HtmlReader()
    documents = reader.load_data(
        Path(__file__).parent / "resources" / "html" / "dummy.html"
    )

    assert len(documents)


def test_pdf_reader():
    reader = AutoReader("PDFReader")
    dirpath = Path(__file__).parent
    documents = reader.load_data(dirpath / "resources" / "dummy.pdf")

    # check document reader output
    assert len(documents) == 1

    first_doc = documents[0]
    assert isinstance(first_doc, Document)
    assert first_doc.text.lower().replace(" ", "") == "dummypdffile"

    langchain_doc = first_doc.to_langchain_format()
    assert isinstance(langchain_doc, LangchainDocument)

    # test chunking using NodeParser from llama-index
    node_parser = SimpleNodeParser.from_defaults(chunk_size=100, chunk_overlap=20)
    nodes = node_parser.get_nodes_from_documents(documents)
    assert len(nodes) > 0


@skip_when_unstructured_pdf_not_installed
def test_unstructured_pdf_reader():
    reader = UnstructuredReader()
    dirpath = Path(__file__).parent
    input_path = dirpath / "resources/dummy.pdf"
    documents = reader.load_data(input_path)

    # check document reader output
    assert len(documents) == 1

    first_doc = documents[0]
    assert isinstance(first_doc, Document)
    assert first_doc.text.lower().replace(" ", "") == "dummypdffile"

    # split documents mode
    documents = reader.load_data(input_path, split_documents=True)
    # check document reader output
    assert len(documents) == 1


def test_mhtml_reader():
    reader = MhtmlReader()
    input_path = Path(__file__).parent / "resources" / "dummy.mhtml"
    docs = reader.load_data(input_path)

    assert len(docs) == 1
    assert docs[0].text.startswith("This is a test")


@patch("azure.ai.documentintelligence.DocumentIntelligenceClient")
def test_azureai_document_intelligence_reader(mock_client):
    reader = AzureAIDocumentIntelligenceLoader(
        endpoint="https://endpoint.com",
        credential="credential",
    )
    docs = reader(Path(__file__).parent / "resources" / "dummy.pdf")

    assert len(docs) == 1
    mock_client.assert_called_once()
