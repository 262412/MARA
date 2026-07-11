from __future__ import annotations

import base64
import hashlib
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMA"
    "ASsJTYQAAAAASUVORK5CYII="
)


def add_hyperlink(paragraph, text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_list_level(paragraph, level: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    numbering = properties.find(qn("w:numPr"))
    if numbering is None:
        numbering = OxmlElement("w:numPr")
        properties.append(numbering)
    indentation = OxmlElement("w:ilvl")
    indentation.set(qn("w:val"), str(level))
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), "999")
    numbering.append(indentation)
    numbering.append(number_id)


def add_picture_with_alt(document: Document, alt_text: str) -> None:
    run = document.add_paragraph().add_run()
    run.add_picture(io.BytesIO(PNG_1X1))
    drawing_properties = run._r.xpath(".//wp:docPr")
    assert len(drawing_properties) == 1
    drawing_properties[0].set("descr", alt_text)


def write_document(path: Path, build) -> Path:
    document = Document()
    build(document)
    document.save(path)
    return path


def replace_image_payload(
    path: Path,
    payload: bytes,
    *,
    content_type: str | None = None,
) -> Path:
    with zipfile.ZipFile(path) as archive:
        members = [(item, archive.read(item.filename)) for item in archive.infolist()]
    image_names = [
        item.filename for item, _payload in members if "/media/" in item.filename
    ]
    assert len(image_names) == 1
    image_name = image_names[0]

    replacement = path.with_suffix(".rewritten.docx")
    with zipfile.ZipFile(replacement, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for item, original_payload in members:
            member_payload = original_payload
            if item.filename == image_name:
                member_payload = payload
            elif item.filename == "[Content_Types].xml" and content_type:
                member_payload = original_payload.replace(
                    b"image/png", content_type.encode("ascii")
                )
            archive.writestr(item, member_payload)
    replacement.replace(path)
    return path


def replace_archive_member(path: Path, member_name: str, payload: bytes) -> Path:
    with zipfile.ZipFile(path) as archive:
        members = [(item, archive.read(item.filename)) for item in archive.infolist()]
    assert member_name in {item.filename for item, _payload in members}

    replacement = path.with_suffix(".rewritten.docx")
    with zipfile.ZipFile(replacement, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for item, original_payload in members:
            archive.writestr(
                item,
                payload if item.filename == member_name else original_payload,
            )
    replacement.replace(path)
    return path


def add_high_ratio_archive_member(path: Path, member_name: str) -> Path:
    with zipfile.ZipFile(path) as archive:
        members = [(item, archive.read(item.filename)) for item in archive.infolist()]

    replacement = path.with_suffix(".rewritten.docx")
    with zipfile.ZipFile(replacement, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for item, payload in members:
            archive.writestr(item, payload)
        archive.writestr(member_name, bytes(4 * 1024 * 1024))
    replacement.replace(path)

    with zipfile.ZipFile(path) as archive:
        info = archive.getinfo(member_name)
    assert info.file_size / info.compress_size > 1_000
    return path


def invalidate_first_table_grid_span(path: Path) -> Path:
    with zipfile.ZipFile(path) as archive:
        payload = archive.read("word/document.xml")
    assert b"<w:tcPr>" in payload
    malformed = payload.replace(
        b"<w:tcPr>",
        b"<w:tcPr><w:gridSpan/>",
        1,
    )
    return replace_archive_member(path, "word/document.xml", malformed)


def incompressible_text(size: int) -> str:
    blocks = (
        hashlib.sha256(str(index).encode("ascii")).hexdigest()
        for index in range((size + 63) // 64)
    )
    return "".join(blocks)[:size]
