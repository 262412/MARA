from __future__ import annotations

import logging

import pytest
from docx.shared import Pt, RGBColor

from .docx_preview_test_utils import (
    PNG_1X1,
    add_high_ratio_archive_member,
    add_hyperlink,
    add_picture_with_alt,
    incompressible_text,
    invalidate_first_table_grid_span,
    replace_image_payload,
    set_list_level,
    write_document,
)


@pytest.fixture(autouse=True)
def _temporary_app_data(monkeypatch, tmp_path):
    monkeypatch.setenv("KH_APP_DATA_DIR", str(tmp_path / "app-data"))


def test_rich_runs_and_safe_hyperlink_keep_existing_html(tmp_path):
    def build(document):
        document.add_heading("Quarterly Plan", level=1)
        paragraph = document.add_paragraph()
        paragraph.add_run("Revenue ")
        bold = paragraph.add_run("grew")
        bold.bold = True
        italic = paragraph.add_run(" quickly")
        italic.italic = True
        underlined = paragraph.add_run(" today")
        underlined.underline = True
        colored = paragraph.add_run(" in blue")
        colored.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        link_paragraph = document.add_paragraph("Read ")
        add_hyperlink(
            link_paragraph, "the report", "https://example.test/report?q=1&x=2"
        )

    source = write_document(tmp_path / "rich.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert html.startswith(
        "<div class='docx-preview' "
        "style=\"font-family:'Times New Roman',serif;font-size:1.00em;\">"
    )
    assert "<h1><span>Quarterly Plan</span></h1>" in html
    assert "<span>Revenue </span>" in html
    assert "<span><strong>grew</strong></span>" in html
    assert "<span><em> quickly</em></span>" in html
    assert "<span><u> today</u></span>" in html
    assert '<span style="color:#123456;"> in blue</span>' in html
    assert (
        '<a href="https://example.test/report?q=1&amp;x=2" target="_blank" '
        'rel="noopener noreferrer"><span>the report</span></a>'
    ) in html


def test_nested_list_markup_is_characterized(tmp_path):
    def build(document):
        parent = document.add_paragraph("Parent")
        set_list_level(parent, 0)
        child = document.add_paragraph("Child")
        set_list_level(child, 1)
        sibling = document.add_paragraph("Sibling")
        set_list_level(sibling, 0)
        document.add_paragraph("After")

    source = write_document(tmp_path / "nested-list.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert (
        "<ul><li><span>Parent</span></li><ul><li><span>Child</span></li>"
        "</ul><li><span>Sibling</span></li></ul><p><span>After</span></p>"
    ) in html


def test_list_first_line_indent_remains_ignored(tmp_path):
    def build(document):
        item = document.add_paragraph("Indented list item")
        item.paragraph_format.first_line_indent = Pt(24)
        set_list_level(item, 0)

    source = write_document(tmp_path / "list-indent.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "<li><span>Indented list item</span></li>" in html
    assert "text-indent" not in html


def test_html_max_chars_keeps_only_complete_existing_paragraphs(tmp_path):
    source = write_document(
        tmp_path / "max-chars.docx",
        lambda document: (
            document.add_paragraph("alpha"),
            document.add_paragraph("bravo"),
        ),
    )

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source), max_chars=5)

    assert "<p><span>alpha</span></p>" in html
    assert "bravo" not in html


def test_paginate_docx_html_keeps_existing_break_and_wrapper_behavior():
    from ktem.pages.chat.page_preview_document import paginate_docx_html

    rich_html = (
        "<div class='docx-preview'><p>First</p>"
        "<p>Second<span class='docx-page-break'></span></p>"
        "<p>Third</p></div>"
    )

    assert paginate_docx_html(rich_html) == [
        "<div class='docx-preview'><p>First</p></div>",
        "<div class='docx-preview'><p>Second<span "
        "class='docx-page-break'></span></p></div>",
        "<div class='docx-preview'><p>Third</p></div>",
    ]
    assert paginate_docx_html("<p>unwrapped</p>") == ["<p>unwrapped</p>"]
    assert paginate_docx_html("") == []


def test_xml_text_is_escaped_before_html_output(tmp_path):
    attack = '<img src=x onerror="alert(1)"> & <script>boom</script>'
    source = write_document(
        tmp_path / "xml-text.docx",
        lambda document: document.add_paragraph(attack),
    )

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert attack not in html
    assert "&lt;img src=x onerror=&quot;alert(1)&quot;&gt;" in html
    assert "&lt;script&gt;boom&lt;/script&gt;" in html


@pytest.mark.parametrize(
    "target",
    [
        "javascript:alert(document.domain)",
        "data:text/html,<script>alert(1)</script>",
        "file:///etc/passwd",
        "../../../../etc/passwd",
        "java\nscript:alert(1)",
    ],
)
def test_unsafe_hyperlink_targets_render_as_text(tmp_path, target):
    def build(document):
        paragraph = document.add_paragraph()
        add_hyperlink(paragraph, "Open safely", target)

    source = write_document(tmp_path / "unsafe-link.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "Open safely" in html
    assert "<a " not in html
    assert target not in html


@pytest.mark.parametrize(
    "target",
    [
        "http://example.test/report",
        "https://example.test/report",
        "mailto:author@example.test",
    ],
)
def test_allowed_hyperlink_targets_keep_existing_anchor_output(tmp_path, target):
    def build(document):
        paragraph = document.add_paragraph()
        add_hyperlink(paragraph, "Open report", target)

    source = write_document(tmp_path / "safe-link.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert f'<a href="{target}" target="_blank"' in html
    assert "<span>Open report</span></a>" in html


def test_table_content_is_rendered_in_document_order_and_escaped(tmp_path):
    def build(document):
        document.add_paragraph("Before")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "<script>alert(1)</script>"
        document.add_paragraph("After")

    source = write_document(tmp_path / "table.docx", build)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "<table>" in html
    assert "<td><p><span>Metric</span></p></td>" in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
    assert html.index("Before") < html.index("<table>") < html.index("After")


def test_safe_raster_image_and_alt_text_are_rendered_safely(tmp_path):
    attack_alt = '"><script>alert("alt")</script>'
    source = write_document(
        tmp_path / "image.docx",
        lambda document: add_picture_with_alt(document, attack_alt),
    )

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "data:image/png;base64," in html
    assert attack_alt not in html
    assert "&quot;&gt;&lt;script&gt;alert(&quot;" in html
    assert "<script>" not in html


@pytest.mark.parametrize(
    ("payload", "content_type"),
    [
        (b"<html><script>alert(1)</script></html>", "text/html"),
        (b"<svg xmlns='http://www.w3.org/2000/svg'><script/></svg>", "image/svg+xml"),
        (b"\xff\xd8\xff\xe0not-a-png", None),
    ],
)
def test_script_capable_or_mismatched_images_never_enter_data_urls(
    tmp_path,
    payload,
    content_type,
):
    source = write_document(
        tmp_path / "unsafe-image.docx",
        lambda document: add_picture_with_alt(document, "Unsafe <image>"),
    )
    replace_image_payload(source, payload, content_type=content_type)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "data:" not in html
    assert "Unsafe &lt;image&gt;" in html
    assert "<script" not in html


def test_oversized_decoded_image_never_enters_a_data_url(tmp_path):
    source = write_document(
        tmp_path / "oversized-image.docx",
        lambda document: add_picture_with_alt(document, "Oversized image"),
    )
    oversized_png = PNG_1X1[:8] + bytes(5 * 1024 * 1024)
    replace_image_payload(source, oversized_png)

    from ktem.pages.chat.page_preview_document import extract_docx_html

    html = extract_docx_html(str(source))

    assert "data:image/png" not in html
    assert "Oversized image" in html


def test_image_occurrence_budget_applies_when_max_chars_has_no_image_text(tmp_path):
    def build(document):
        for index in range(20):
            add_picture_with_alt(document, f"Occurrence image {index}")

    source = write_document(tmp_path / "image-count.docx", build)

    from ktem.preview.docx import extract_docx_html_strict

    html = extract_docx_html_strict(str(source), max_chars=1)

    assert html.count("data:image/png;base64,") == 16
    assert "Occurrence image 19" in html
    assert html.endswith("</div>")


def test_aggregate_image_bytes_are_bounded_per_render(tmp_path):
    def build(document):
        for index in range(3):
            add_picture_with_alt(document, f"Aggregate image {index}")

    source = write_document(tmp_path / "image-bytes.docx", build)
    payload = PNG_1X1[:8] + bytes(2 * 1024 * 1024 - 8)
    replace_image_payload(source, payload)

    from ktem.preview.docx import extract_docx_html_strict

    html = extract_docx_html_strict(str(source), max_chars=1)

    assert html.count("data:image/png;base64,") == 2
    assert "Aggregate image 2" in html
    assert html.endswith("</div>")


def test_final_html_budget_keeps_image_only_render_complete(tmp_path):
    oversized_alt = incompressible_text(8 * 1024 * 1024 + 1024)
    source = write_document(
        tmp_path / "html-budget.docx",
        lambda document: add_picture_with_alt(document, oversized_alt),
    )

    from ktem.preview.docx import extract_docx_html_strict

    html = extract_docx_html_strict(str(source), max_chars=1)

    assert len(html) <= 8 * 1024 * 1024
    assert html.startswith("<div class='docx-preview'")
    assert html.endswith("</div>")


def test_table_survives_docx_pagination():
    from ktem.pages.chat.page_preview_document import paginate_docx_html

    rich_html = (
        "<div class='docx-preview'><p>Before</p><table><tbody><tr>"
        "<td>Value</td></tr></tbody></table><p>After</p></div>"
    )

    pages = paginate_docx_html(rich_html)

    assert len(pages) == 1
    assert "<table><tbody><tr><td>Value</td></tr></tbody></table>" in pages[0]


def test_strict_corrupt_docx_error_has_path_stage_and_converter(tmp_path):
    from ktem.preview.docx import extract_docx_html_strict
    from ktem.preview.errors import PreviewErrorCode, PreviewSourceError

    corrupt = tmp_path / "truncated.docx"
    corrupt.write_bytes(b"PK\x03\x04truncated")

    with pytest.raises(PreviewSourceError) as caught:
        extract_docx_html_strict(str(corrupt))

    assert caught.value.code is PreviewErrorCode.SOURCE_ARCHIVE_INVALID
    assert caught.value.stage == "docx_package"
    assert caught.value.source_path == corrupt.resolve()
    assert caught.value.converter == "python-docx"
    assert "DOCX" in caught.value.details


@pytest.mark.parametrize(
    "member_name",
    [
        "word/review-bomb.bin",
        "word/_rels/review-bomb.rels",
    ],
)
@pytest.mark.parametrize(
    "strict_name", ["extract_docx_text_strict", "extract_docx_html_strict"]
)
def test_strict_docx_rejects_high_ratio_entries_before_package_parsing(
    tmp_path, strict_name, member_name
):
    import ktem.preview.docx as docx_preview
    from ktem.preview.errors import PreviewErrorCode, PreviewSourceError

    source = write_document(
        tmp_path / f"high-ratio-{strict_name}.docx",
        lambda document: document.add_paragraph("bounded"),
    )
    add_high_ratio_archive_member(source, member_name)
    strict_extract = getattr(docx_preview, strict_name)

    with pytest.raises(PreviewSourceError) as caught:
        strict_extract(str(source))

    assert caught.value.code is PreviewErrorCode.SOURCE_ARCHIVE_INVALID
    assert caught.value.stage == "archive_validation"
    assert caught.value.source_path == source.resolve()
    assert caught.value.converter == "python-docx"
    assert "compression ratio" in caught.value.details.lower()


def test_invalid_required_table_xml_raises_typed_render_error(tmp_path):
    from ktem.preview.docx import extract_docx_html_strict
    from ktem.preview.errors import PreviewErrorCode, PreviewSourceError

    source = write_document(
        tmp_path / "invalid-table.docx",
        lambda document: document.add_table(rows=1, cols=1),
    )
    invalidate_first_table_grid_span(source)

    with pytest.raises(PreviewSourceError) as caught:
        extract_docx_html_strict(str(source))

    assert caught.value.code is PreviewErrorCode.SOURCE_INVALID
    assert caught.value.stage == "docx_render"
    assert caught.value.source_path == source.resolve()
    assert caught.value.converter == "python-docx"
    assert "malformed" in caught.value.details.lower()


@pytest.mark.parametrize(
    "strict_name", ["extract_docx_text_strict", "extract_docx_html_strict"]
)
@pytest.mark.parametrize("invalid_source", [None, "invalid\0source.docx"])
def test_strict_invalid_docx_paths_raise_typed_source_errors(
    strict_name, invalid_source
):
    import ktem.preview.docx as docx_preview
    from ktem.preview.errors import PreviewErrorCode, PreviewSourceError

    strict_extract = getattr(docx_preview, strict_name)

    with pytest.raises(PreviewSourceError) as caught:
        strict_extract(invalid_source)

    assert caught.value.code is PreviewErrorCode.SOURCE_INVALID
    assert caught.value.stage == "docx_source"
    assert caught.value.converter == "python-docx"
    assert repr(invalid_source) in caught.value.details


def test_compat_corrupt_docx_fallback_logs_actionable_context(tmp_path, caplog):
    from ktem.pages.chat.page_preview_document import extract_docx_html

    corrupt = tmp_path / "bad-<document>.docx"
    corrupt.write_bytes(b"not a zip archive")

    with caplog.at_level(logging.WARNING):
        html = extract_docx_html(str(corrupt))

    assert html == ""
    assert str(corrupt.resolve()) in caplog.text
    assert "stage=docx_package" in caplog.text
    assert "converter=python-docx" in caplog.text
