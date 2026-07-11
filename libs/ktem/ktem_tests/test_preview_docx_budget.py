from __future__ import annotations

from html import escape as html_escape

import pytest
from docx import Document
from docx.oxml.ns import qn

from .docx_preview_test_utils import add_hyperlink, add_picture_with_alt


def _render(document: Document, *, max_chars: int) -> str:
    from ktem.preview.docx_render import DocxHtmlRenderer

    return DocxHtmlRenderer(document).render(max_chars=max_chars)


def test_remaining_html_budget_rejects_image_before_base64(monkeypatch):
    from ktem.preview import docx_security
    from ktem.preview.docx_security import MAX_RENDERED_HTML_CHARS

    document = Document()
    document.add_paragraph("a" * (MAX_RENDERED_HTML_CHARS - 250))
    add_picture_with_alt(document, "Budget image")
    calls = 0
    original_encode = docx_security.base64.b64encode

    def recording_encode(payload):
        nonlocal calls
        calls += 1
        return original_encode(payload)

    monkeypatch.setattr(docx_security.base64, "b64encode", recording_encode)

    html = _render(document, max_chars=MAX_RENDERED_HTML_CHARS)

    assert calls == 0
    assert len(html) <= MAX_RENDERED_HTML_CHARS
    assert html.endswith("</div>")


@pytest.mark.parametrize(
    ("hostile_char", "repeat_count"),
    [('"', 2 * 1024 * 1024), ("'", 1_500_000)],
)
def test_huge_text_is_rejected_before_escape_allocation(
    monkeypatch,
    hostile_char,
    repeat_count,
):
    from ktem.preview import docx_runs
    from ktem.preview.docx_security import MAX_RENDERED_HTML_CHARS

    hostile_text = hostile_char * repeat_count
    document = Document()
    document.add_paragraph(hostile_text)

    def guarded_escape(value, quote=True):
        if value == hostile_text:
            pytest.fail("oversized run text reached html.escape")
        return html_escape(value, quote=quote)

    monkeypatch.setattr(docx_runs, "escape", guarded_escape)

    html = _render(document, max_chars=len(hostile_text))

    assert len(html) <= MAX_RENDERED_HTML_CHARS
    assert hostile_text not in html
    assert html.endswith("</div>")


@pytest.mark.parametrize(
    ("hostile_char", "repeat_count"),
    [('"', 2 * 1024 * 1024), ("'", 1_500_000)],
)
def test_huge_hyperlink_is_rejected_before_escape_allocation(
    monkeypatch,
    hostile_char,
    repeat_count,
):
    from ktem.preview import docx_runs
    from ktem.preview.docx_security import MAX_RENDERED_HTML_CHARS

    hostile_target = "https://example.test/" + hostile_char * repeat_count
    document = Document()
    paragraph = document.add_paragraph()
    add_hyperlink(paragraph, "Safe label", hostile_target)

    def guarded_escape(value, quote=True):
        if value == hostile_target:
            pytest.fail("oversized hyperlink reached html.escape")
        return html_escape(value, quote=quote)

    monkeypatch.setattr(docx_runs, "escape", guarded_escape)

    html = _render(document, max_chars=len("Safe label"))

    assert len(html) <= MAX_RENDERED_HTML_CHARS
    assert "Safe label" in html
    assert "<a " not in html


@pytest.mark.parametrize(
    ("hostile_char", "repeat_count"),
    [('"', 2 * 1024 * 1024), ("'", 1_500_000)],
)
def test_huge_image_alt_is_rejected_before_escape_allocation(
    monkeypatch,
    hostile_char,
    repeat_count,
):
    from ktem.preview import docx_runs
    from ktem.preview.docx_security import MAX_RENDERED_HTML_CHARS

    hostile_alt = hostile_char * repeat_count
    document = Document()
    add_picture_with_alt(document, hostile_alt)
    blip = document.element.body.xpath(".//a:blip")[0]
    blip.set(qn("r:embed"), "missing-relationship")

    def guarded_escape(value, quote=True):
        if value == hostile_alt:
            pytest.fail("oversized image alt reached html.escape")
        return html_escape(value, quote=quote)

    monkeypatch.setattr(docx_runs, "escape", guarded_escape)

    html = _render(document, max_chars=1)

    assert len(html) <= MAX_RENDERED_HTML_CHARS
    assert hostile_alt not in html
    assert html.endswith("</div>")


def test_single_quote_escape_length_matches_html_escape():
    from ktem.preview.docx_security import escaped_html_length

    value = "'" * 4096

    assert escaped_html_length(value) == len(html_escape(value))
